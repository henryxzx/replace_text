import docx
import os
import xlrd


def info_update(doc, old_info, new_info):
    """
    此函数用于批量替换文档中需要替换的信息
    doc:文件
    old_info和new_info：原文字和需要替换的新文字
    """

    # 读取段落中的所有run，找到需替换的信息进行替换
    for p in doc.paragraphs:
        if old_info in p.text:
            inline = p.runs
            for i in inline:
                if old_info in i.text:
                    text = i.text.replace(old_info, new_info)
                    i.text = text

    # 读取表格中的所有单元格，找到需替换的信息进行替换
    for table in doc.tables:
        for r in table.rows:
            mycell = r.cells
            for c in mycell:
                for p in c.paragraphs:
                    if old_info in p.text:
                        inline = p.runs
                        for i in inline:
                            if old_info in i.text:
                                text = i.text.replace(old_info, new_info)
                                i.text = text


def folder_exists(path):
    tar_path = os.path.join(path, '替换结果')
    is_exists = os.path.exists(tar_path)
    if not is_exists:
        print('替换结果目录不存在！创建新的替换结果目录')
        if os.makedirs(tar_path):
            print(path + ' 创建成功')
    else:
        print('替换结果目录已存在！替换后的文档将保存在目录下')


def info_read_from_excel():
    path = os.getcwd() # 文件夹路径
    folder_exists(path)  # 判断替换内容文件夹是否存在当前目录下
    excel_file = os.path.join(path, '替换内容.xlsx')
    if not os.path.isfile(excel_file):
        print('替换内容.xlsx 不存在，请在目录下创建')
        quit()
    else:
        print('正在读取替换内容.xlsx...')
    wb = xlrd.open_workbook(filename=excel_file)  # 打开文件
    sheet1 = wb.sheet_by_index(0)  # 通过索引获取表格
    old_info = []
    new_info = []
    for i in range(sheet1.nrows):
        old_info.append(sheet1.cell(i, 0).value)
        new_info.append(sheet1.cell(i, 1).value)

    files = []  # 目录下的文件
    for file in os.listdir(path):
        if file.endswith(".docx"):  # 排除文件夹内的其它干扰文件，只获取word文件
            files.append(os.path.join(path, file))

    for file in files:
        doc = docx.Document(file)
        for n in range(0, len(old_info)):
            info_update(doc, old_info[n], new_info[n])
        doc.save("替换结果/{}".format(file.split("/")[-1]))
        # doc.save("替换结果/{}".format(file.split("\\")[-1]))
        print("{} 替换完成".format(file))


if __name__ == '__main__':
    info_read_from_excel()
    print('运行成功，请输入任意键退出程序......')
    input()
